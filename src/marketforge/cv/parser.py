"""
MarketForge AI — CV Parser.

Extracts structured text from PDF and DOCX files entirely in-memory.
No file is written to disk at any point.

Outputs a ParsedCV dataclass with:
  - raw_text:  full extracted text (PII not yet stripped — strip before LLM)
  - sections:  dict keyed by section name (experience, education, skills, …)
  - metadata:  page_count, has_tables, has_images, estimated_years_experience
  - flags:     has_email, has_phone, has_linkedin, has_github

PDF strategy:   pdfplumber (layout-aware) → pypdf (text-only fallback)
DOCX strategy:  python-docx
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

import structlog

logger = structlog.get_logger(__name__)

# ── Section header detection ──────────────────────────────────────────────────
# Patterns use a full-line anchor (\s*$) so that body text starting with a
# keyword (e.g. "Experienced ML engineer") is NOT mistaken for a header.
# An optional trailing colon or dash is allowed (e.g. "Experience:" / "Skills —").
_SECTION_PATTERNS: dict[str, re.Pattern] = {
    "summary":        re.compile(r"^(summary|profile|objective|about\s*me|professional\s*summary|personal\s*statement)\s*[:\-—]?\s*$", re.I),
    "experience":     re.compile(r"^(experience|work\s*experience|employment|career\s*history|work\s*history|professional\s*experience)\s*[:\-—]?\s*$", re.I),
    "education":      re.compile(r"^(education|academic|qualifications?|degrees?|university|schooling)\s*[:\-—]?\s*$", re.I),
    "skills":         re.compile(r"^(skills|technical\s*skills|core\s*competencies|expertise|technologies|key\s*skills)\s*[:\-—]?\s*$", re.I),
    "certifications": re.compile(r"^(certifications?|licen[sc]es?|awards?|achievements?|accreditation)\s*[:\-—]?\s*$", re.I),
    "contact":        re.compile(r"^(contact|personal\s*details?|personal\s*information)\s*[:\-—]?\s*$", re.I),
    "projects":       re.compile(r"^(projects?|portfolio|open\s*source|side\s*projects?)\s*[:\-—]?\s*$", re.I),
    "publications":   re.compile(r"^(publications?|research|papers?|patents?)\s*[:\-—]?\s*$", re.I),
}

# ── PII presence flags (for metadata — actual scrubbing done in gdpr.py) ──────
_EMAIL_RE    = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
_PHONE_RE    = re.compile(r"\b(?:\+44|0)[\s\-]?\d{4}[\s\-]?\d{6}\b")
_LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.I)
_GITHUB_RE   = re.compile(r"github\.com/[\w\-]+", re.I)
_YEAR_RE     = re.compile(r"\b(20[0-2]\d|199\d)\b")


@dataclass
class ParsedCV:
    raw_text:        str
    sections:        dict[str, str] = field(default_factory=dict)
    has_email:       bool = False
    has_phone:       bool = False
    has_linkedin:    bool = False
    has_github:      bool = False
    page_count:      int  = 1
    has_tables:      bool = False
    has_images:      bool = False
    estimated_years: float | None = None
    parse_method:    str  = "unknown"
    error:           str | None = None


def parse_cv(data: bytes, file_type: str) -> ParsedCV:
    """
    Parse CV bytes into a ParsedCV.
    file_type must be "pdf" or "docx" (as returned by scanner.ScanResult.file_type).
    All processing is in-memory — data bytes are not written to disk.
    """
    if file_type == "pdf":
        return _parse_pdf(data)
    if file_type == "docx":
        return _parse_docx(data)
    return ParsedCV(raw_text="", error="unsupported_file_type")


# ── PDF ────────────────────────────────────────────────────────────────────────

def _parse_pdf(data: bytes) -> ParsedCV:
    """Primary: pdfplumber (layout-aware, detects tables/images)."""
    try:
        import pdfplumber
        pages:      list[str] = []
        has_tables: bool      = False
        has_images: bool      = False

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                if page.extract_tables():
                    has_tables = True
                if page.images:
                    has_images = True
                text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                pages.append(text)

        raw = "\n".join(pages).strip()
        return _build(raw, page_count, has_tables, has_images, "pdfplumber")

    except Exception as exc:
        logger.warning("cv.parse.pdfplumber_failed", error=str(exc))
        return _parse_pdf_fallback(data)


def _parse_pdf_fallback(data: bytes) -> ParsedCV:
    """Fallback: pypdf (pure text, no layout info)."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages  = [page.extract_text() or "" for page in reader.pages]
        raw    = "\n".join(pages).strip()
        return _build(raw, len(reader.pages), False, False, "pypdf_fallback")
    except Exception as exc:
        logger.error("cv.parse.pdf_failed", error=str(exc))
        return ParsedCV(raw_text="", error="pdf_parse_failed")


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _parse_docx(data: bytes) -> ParsedCV:
    """python-docx: extract paragraphs + table cells."""
    try:
        from docx import Document
        doc        = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        has_tables = len(doc.tables) > 0

        # Include table cell text so skills in tables aren't missed entirely
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        raw = "\n".join(paragraphs).strip()
        return _build(raw, 1, has_tables, False, "python_docx")

    except Exception as exc:
        logger.error("cv.parse.docx_failed", error=str(exc))
        return ParsedCV(raw_text="", error="docx_parse_failed")


# ── Shared builder ─────────────────────────────────────────────────────────────

def _build(
    raw:        str,
    page_count: int,
    has_tables: bool,
    has_images: bool,
    method:     str,
) -> ParsedCV:
    return ParsedCV(
        raw_text        = raw,
        sections        = _detect_sections(raw),
        has_email       = bool(_EMAIL_RE.search(raw)),
        has_phone       = bool(_PHONE_RE.search(raw)),
        has_linkedin    = bool(_LINKEDIN_RE.search(raw)),
        has_github      = bool(_GITHUB_RE.search(raw)),
        page_count      = page_count,
        has_tables      = has_tables,
        has_images      = has_images,
        estimated_years = _estimate_years(raw),
        parse_method    = method,
    )


def _detect_sections(text: str) -> dict[str, str]:
    """
    Split raw CV text into named sections by detecting header lines.
    A line is treated as a section header when it:
      - matches a known header pattern
      - is short enough to be a heading (<= 60 chars)
    """
    lines:   list[str]       = text.split("\n")
    sections: dict[str, str] = {}
    current:  str            = "preamble"
    buf:      list[str]      = []

    for line in lines:
        stripped = line.strip()
        matched: str | None = None

        if stripped and len(stripped) <= 60:
            for section, pattern in _SECTION_PATTERNS.items():
                if pattern.match(stripped):
                    matched = section
                    break

        if matched:
            if buf:
                sections[current] = "\n".join(buf).strip()
            current = matched
            buf     = []
        else:
            buf.append(line)

    if buf:
        sections[current] = "\n".join(buf).strip()

    return {k: v for k, v in sections.items() if v}


def _estimate_years(text: str) -> float | None:
    """
    Crude heuristic: find the earliest year mentioned and compute span to today.
    E.g. "2018 – 2024" with today=2026 → ~8 years career span.
    """
    from datetime import date
    years = [int(y) for y in _YEAR_RE.findall(text)]
    if not years:
        return None
    current_year = date.today().year
    return max(0.0, float(current_year - min(years)))
